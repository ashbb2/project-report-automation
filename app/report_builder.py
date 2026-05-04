import re
import time
from urllib.parse import urlparse
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from typing import Dict, Any, List
from io import BytesIO
from app.llm_client import llm_client
from app.config import Config
from app.prompt_renderer import get_section_prompt
from app.db import get_cached_section, save_section, upsert_report_status

SECTION_LABELS = {
    'executive_summary':       'Executive Summary',
    'introduction':            'Introduction',
    'regulatory_framework':    'Regulatory Framework',
    'market_assessment':       'Market Assessment',
    'business_operating_model':'Business & Operating Model',
    'equipment_profiles':      'Equipment Profiles',
    'financial_feasibility':   'Financial Feasibility',
    'risk_assessment':         'Risk Assessment',
    'caveats':                 'Caveats',
    'appendices':              'Appendices',
}


def identify_missing_inputs(submission: Dict[str, Any]) -> List[str]:
    """
    Identify which fields are missing or blank in the submission.
    
    Args:
        submission: Dictionary containing submission data
        
    Returns:
        List of missing field names
    """
    required_fields = [
        'business_idea', 'location_land', 'promoter_background',
        'goals', 'start_date', 'target_launch_date', 'budget', 'target_market'
    ]
    
    missing = []
    for field in required_fields:
        value = submission.get(field)
        if not value or (isinstance(value, str) and not value.strip()):
            missing.append(field)
    
    return missing


def apply_report_formatting(doc: Document) -> None:
    """
    Apply the report-wide typography and spacing specification.
    """
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(6)
    paragraph_format.space_after = Pt(6)
    paragraph_format.line_spacing = 1.0


# ---------------------------------------------------------------------------
# Markdown → Word helpers
# ---------------------------------------------------------------------------

URL_PATTERN = re.compile(r"(https?://[^\s)]+)")
MD_LINK_PATTERN = re.compile(r"\[([^\]]+)\]\((https?://[^\s)]+)\)")


def _is_valid_url_syntax(url: str) -> bool:
    try:
        parsed = urlparse(url)
        return parsed.scheme in {"http", "https"} and bool(parsed.netloc)
    except Exception:
        return False


def _sanitize_invalid_links(text: str, validation_cache: Dict[str, bool]) -> str:
    """Remove links with invalid syntax; all syntactically-valid URLs are kept as-is.
    Network reachability checks are intentionally skipped — they cause multi-minute
    blocking delays when AI-generated sections contain many URLs.
    """
    if not text:
        return text

    def is_ok(url: str) -> bool:
        if url not in validation_cache:
            validation_cache[url] = _is_valid_url_syntax(url)
        return validation_cache[url]

    def md_link_repl(match):
        label, url = match.group(1), match.group(2)
        if is_ok(url):
            return match.group(0)
        return f"{label} (Source link unavailable)"

    updated = MD_LINK_PATTERN.sub(md_link_repl, text)

    def bare_url_repl(match):
        url = match.group(1)
        if is_ok(url):
            return url
        return "Source link unavailable"

    return URL_PATTERN.sub(bare_url_repl, updated)


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    """Insert a clickable external hyperlink into a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)

    new_run.append(rPr)
    text_el = OxmlElement('w:t')
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_inline_runs(paragraph, text: str) -> None:
    """Render inline markdown features: **bold**, [links](url), and bare URLs."""
    # Split by bold first; odd indexes are bold segments.
    bold_parts = re.split(r'\*\*(.+?)\*\*', text)
    for i, part in enumerate(bold_parts):
        if not part:
            continue

        is_bold = (i % 2 == 1)
        cursor = 0

        # First resolve markdown links [text](url)
        for m in MD_LINK_PATTERN.finditer(part):
            before = part[cursor:m.start()]
            if before:
                # Handle bare URLs in the non-link segment
                start = 0
                for u in URL_PATTERN.finditer(before):
                    plain = before[start:u.start()]
                    if plain:
                        run = paragraph.add_run(plain)
                        run.bold = is_bold
                    _add_hyperlink(paragraph, u.group(1), u.group(1))
                    start = u.end()
                tail = before[start:]
                if tail:
                    run = paragraph.add_run(tail)
                    run.bold = is_bold

            _add_hyperlink(paragraph, m.group(1), m.group(2))
            cursor = m.end()

        remaining = part[cursor:]
        if remaining:
            start = 0
            for u in URL_PATTERN.finditer(remaining):
                plain = remaining[start:u.start()]
                if plain:
                    run = paragraph.add_run(plain)
                    run.bold = is_bold
                _add_hyperlink(paragraph, u.group(1), u.group(1))
                start = u.end()
            tail = remaining[start:]
            if tail:
                run = paragraph.add_run(tail)
                run.bold = is_bold


def _split_md_table_row(line: str) -> List[str]:
    row = line.strip()
    if row.startswith('|'):
        row = row[1:]
    if row.endswith('|'):
        row = row[:-1]
    return [c.strip() for c in row.split('|')]


def _is_md_table_separator(line: str) -> bool:
    s = line.strip()
    return bool(re.match(r'^\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?$', s))

def _add_para_with_inline_bold(doc: Document, text: str, style: str = None):
    """Add a paragraph, converting **bold** spans to Word bold runs."""
    para = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    _add_inline_runs(para, text)
    return para


def render_markdown_to_doc(doc: Document, text: str) -> None:
    """
    Parse LLM markdown output and write it into the Word document
    using proper Word styles instead of raw markdown symbols.
    """
    if not text:
        return

    lines = text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        # Markdown table: header row + separator + data rows
        if '|' in stripped and i + 1 < len(lines) and _is_md_table_separator(lines[i + 1]):
            header_cells = _split_md_table_row(lines[i])
            table_rows = []
            j = i + 2
            while j < len(lines):
                row_line = lines[j].strip()
                if not row_line or '|' not in row_line:
                    break
                # Skip accidental separator repeats
                if _is_md_table_separator(row_line):
                    j += 1
                    continue
                table_rows.append(_split_md_table_row(lines[j]))
                j += 1

            col_count = max(1, len(header_cells))
            table = doc.add_table(rows=1 + len(table_rows), cols=col_count)
            table.style = 'Table Grid'

            for c, value in enumerate(header_cells[:col_count]):
                cell_para = table.rows[0].cells[c].paragraphs[0]
                cell_para.text = ''
                _add_inline_runs(cell_para, value)
                for run in cell_para.runs:
                    run.bold = True

            for r, row_vals in enumerate(table_rows, start=1):
                for c in range(col_count):
                    value = row_vals[c] if c < len(row_vals) else ''
                    cell_para = table.rows[r].cells[c].paragraphs[0]
                    cell_para.text = ''
                    _add_inline_runs(cell_para, value)

            i = j
            continue

        if stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=2)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            _add_para_with_inline_bold(doc, stripped[2:], style='List Bullet')
        elif stripped.startswith('**') and stripped.endswith('**') and stripped.count('**') == 2:
            p = doc.add_paragraph()
            p.add_run(stripped[2:-2]).bold = True
        else:
            _add_para_with_inline_bold(doc, stripped)
        i += 1


# ---------------------------------------------------------------------------
# Table of Contents
# ---------------------------------------------------------------------------

def _add_toc(doc: Document) -> None:
    """Insert a TOC field. Word will populate it on first open (Ctrl+A, F9)."""
    doc.add_heading('Table of Contents', level=1)
    para = doc.add_paragraph()
    run = para.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')
    placeholder = OxmlElement('w:r')
    placeholder_t = OxmlElement('w:t')
    placeholder_t.text = '[Right-click here and select "Update Field" to generate the Table of Contents]'
    placeholder.append(placeholder_t)
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_sep)
    run._r.append(placeholder)
    run._r.append(fldChar_end)
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Financial projection helpers
# ---------------------------------------------------------------------------

def _parse_budget_inr(budget_str: str) -> float:
    """Parse strings like '50 lakhs', '2 crores', '5000000' to a float in INR."""
    s = str(budget_str).lower().replace(',', '').strip()
    for word, mult in [
        ('crores', 1e7), ('crore', 1e7), ('cr', 1e7),
        ('lakhs', 1e5), ('lakh', 1e5), ('lacs', 1e5), ('lac', 1e5),
        ('millions', 1e6), ('million', 1e6), ('mn', 1e6),
    ]:
        if word in s:
            try:
                num_str = s[:s.index(word)].strip()
                num = float(re.sub(r'[^\d.]', '', num_str) or '0')
                if num > 0:
                    return num * mult
            except Exception:
                pass
    try:
        cleaned = re.sub(r'[^\d.]', '', s)
        return float(cleaned) if cleaned else 5_000_000
    except Exception:
        return 5_000_000


def _fmt(val: float) -> str:
    """Format a float as a readable INR string."""
    if val >= 1e7:
        return f'\u20b9{val / 1e7:.2f} Cr'
    elif val >= 1e5:
        return f'\u20b9{val / 1e5:.2f} L'
    else:
        return f'\u20b9{val:,.0f}'


def _compute_financials(budget: float) -> dict:
    """Derive a full 3-year financial model from the total project cost."""
    land         = budget * 0.10
    building     = budget * 0.20
    plant        = budget * 0.40
    misc_fixed   = budget * 0.05
    preop        = budget * 0.05
    wc_margin    = budget * 0.20
    equity       = budget * 0.30
    term_loan    = budget * 0.50
    wc_loan      = budget * 0.20

    rev_full = budget * 2.0
    rev  = {1: rev_full * 0.60, 2: rev_full * 0.75, 3: rev_full * 0.90}
    rm   = {y: rev[y] * 0.40 for y in [1, 2, 3]}
    util = {y: rev[y] * 0.05 for y in [1, 2, 3]}
    labor = {y: rev[y] * 0.12 for y in [1, 2, 3]}
    admin = {y: rev[y] * 0.05 for y in [1, 2, 3]}
    depreciation = plant * 0.10 + building * 0.05
    interest = {1: term_loan * 0.12, 2: term_loan * 0.12 * 0.85, 3: term_loan * 0.12 * 0.70}
    ebitda = {y: rev[y] - rm[y] - util[y] - labor[y] - admin[y] for y in [1, 2, 3]}
    ebit   = {y: ebitda[y] - depreciation for y in [1, 2, 3]}
    pbt    = {y: ebit[y] - interest[y] for y in [1, 2, 3]}
    tax    = {y: max(0.0, pbt[y] * 0.25) for y in [1, 2, 3]}
    pat    = {y: pbt[y] - tax[y] for y in [1, 2, 3]}

    return dict(
        land=land, building=building, plant=plant, misc_fixed=misc_fixed,
        preop=preop, wc_margin=wc_margin, total=budget,
        equity=equity, term_loan=term_loan, wc_loan=wc_loan,
        rev=rev, rm=rm, util=util, labor=labor, admin=admin,
        depreciation=depreciation, interest=interest,
        ebitda=ebitda, ebit=ebit, pbt=pbt, tax=tax, pat=pat,
    )


def _build_table_data(index: int, f: dict):
    """Return (headers, rows) for a financial table by index (1-based)."""
    fmt = _fmt
    principal = f['term_loan'] * 0.15

    if index == 1:
        headers = ['Cost Head', 'Amount', '% of Total']
        rows = [
            ['Land & Site Development',     fmt(f['land']),       '10%'],
            ['Civil Construction & Building', fmt(f['building']),  '20%'],
            ['Plant & Machinery',            fmt(f['plant']),      '40%'],
            ['Misc. Fixed Assets',           fmt(f['misc_fixed']), '5%'],
            ['Pre-operative Expenses',       fmt(f['preop']),      '5%'],
            ['Working Capital Margin',       fmt(f['wc_margin']),  '20%'],
            ['Total Project Cost',           fmt(f['total']),      '100%'],
        ]
    elif index == 2:
        headers = ['Source', 'Amount', '% of Total']
        rows = [
            ['Promoter Equity',      fmt(f['equity']),     '30%'],
            ['Term Loan (Bank/FI)',  fmt(f['term_loan']),  '50%'],
            ['Working Capital Loan', fmt(f['wc_loan']),    '20%'],
            ['Total',                fmt(f['total']),      '100%'],
        ]
    elif index == 3:
        headers = ['Parameter', 'Year 1', 'Year 2', 'Year 3']
        rows = [
            ['Installed Capacity (%)',         '100%', '100%', '100%'],
            ['Utilization (%)',                '60%',  '75%',  '90%'],
            ['Effective Production (% of max)', '60%', '75%',  '90%'],
        ]
    elif index == 4:
        headers = ['Item', 'Year 1', 'Year 2', 'Year 3']
        rows = [
            ['Gross Revenue',                fmt(f['rev'][1]),          fmt(f['rev'][2]),          fmt(f['rev'][3])],
            ['Less: Returns & Discounts (2%)', fmt(f['rev'][1]*0.02),   fmt(f['rev'][2]*0.02),     fmt(f['rev'][3]*0.02)],
            ['Net Revenue',                  fmt(f['rev'][1]*0.98),     fmt(f['rev'][2]*0.98),     fmt(f['rev'][3]*0.98)],
        ]
    elif index == 5:
        headers = ['Item', 'Year 1', 'Year 2', 'Year 3']
        rows = [
            ['Primary Raw Material (70%)',       fmt(f['rm'][1]*0.70), fmt(f['rm'][2]*0.70), fmt(f['rm'][3]*0.70)],
            ['Secondary Raw Material (20%)',     fmt(f['rm'][1]*0.20), fmt(f['rm'][2]*0.20), fmt(f['rm'][3]*0.20)],
            ['Packaging & Consumables (10%)',    fmt(f['rm'][1]*0.10), fmt(f['rm'][2]*0.10), fmt(f['rm'][3]*0.10)],
            ['Total Raw Material Cost',          fmt(f['rm'][1]),      fmt(f['rm'][2]),      fmt(f['rm'][3])],
        ]
    elif index == 6:
        headers = ['Item', 'Year 1', 'Year 2', 'Year 3']
        rows = [
            ['Power & Electricity (60%)',     fmt(f['util'][1]*0.60), fmt(f['util'][2]*0.60), fmt(f['util'][3]*0.60)],
            ['Fuel & Thermal Energy (25%)',   fmt(f['util'][1]*0.25), fmt(f['util'][2]*0.25), fmt(f['util'][3]*0.25)],
            ['Water & Effluent (10%)',        fmt(f['util'][1]*0.10), fmt(f['util'][2]*0.10), fmt(f['util'][3]*0.10)],
            ['Other Utilities (5%)',          fmt(f['util'][1]*0.05), fmt(f['util'][2]*0.05), fmt(f['util'][3]*0.05)],
            ['Total Utility Cost',            fmt(f['util'][1]),      fmt(f['util'][2]),      fmt(f['util'][3])],
        ]
    elif index == 7:
        headers = ['Category', 'Year 1', 'Year 2', 'Year 3']
        rows = [
            ['Production Staff (50%)',    fmt(f['labor'][1]*0.50), fmt(f['labor'][2]*0.50), fmt(f['labor'][3]*0.50)],
            ['Management & Admin (25%)',  fmt(f['labor'][1]*0.25), fmt(f['labor'][2]*0.25), fmt(f['labor'][3]*0.25)],
            ['Sales & Marketing (15%)',   fmt(f['labor'][1]*0.15), fmt(f['labor'][2]*0.15), fmt(f['labor'][3]*0.15)],
            ['Other Staff (10%)',         fmt(f['labor'][1]*0.10), fmt(f['labor'][2]*0.10), fmt(f['labor'][3]*0.10)],
            ['Admin Overhead',            fmt(f['admin'][1]),       fmt(f['admin'][2]),      fmt(f['admin'][3])],
            ['Total Employee & Overhead', fmt(f['labor'][1]+f['admin'][1]), fmt(f['labor'][2]+f['admin'][2]), fmt(f['labor'][3]+f['admin'][3])],
        ]
    elif index == 8:
        dep = f['depreciation']
        headers = ['Line Item', 'Year 1', 'Year 2', 'Year 3']
        rows = [
            ['Net Revenue',                fmt(f['rev'][1]),                    fmt(f['rev'][2]),                    fmt(f['rev'][3])],
            ['Less: Raw Material',         fmt(f['rm'][1]),                     fmt(f['rm'][2]),                     fmt(f['rm'][3])],
            ['Less: Utilities',            fmt(f['util'][1]),                   fmt(f['util'][2]),                   fmt(f['util'][3])],
            ['Less: Employee & Overhead',  fmt(f['labor'][1]+f['admin'][1]),    fmt(f['labor'][2]+f['admin'][2]),    fmt(f['labor'][3]+f['admin'][3])],
            ['EBITDA',                     fmt(f['ebitda'][1]),                  fmt(f['ebitda'][2]),                 fmt(f['ebitda'][3])],
            ['Less: Depreciation',         fmt(dep),                            fmt(dep),                            fmt(dep)],
            ['EBIT',                       fmt(f['ebit'][1]),                   fmt(f['ebit'][2]),                   fmt(f['ebit'][3])],
            ['Less: Interest',             fmt(f['interest'][1]),               fmt(f['interest'][2]),               fmt(f['interest'][3])],
            ['PBT',                        fmt(f['pbt'][1]),                    fmt(f['pbt'][2]),                    fmt(f['pbt'][3])],
            ['Less: Tax @ 25%',            fmt(f['tax'][1]),                    fmt(f['tax'][2]),                    fmt(f['tax'][3])],
            ['PAT (Net Profit)',            fmt(f['pat'][1]),                    fmt(f['pat'][2]),                    fmt(f['pat'][3])],
        ]
    elif index == 9:
        dep = f['depreciation']
        headers = ['Item', 'Year 1', 'Year 2', 'Year 3']
        rows = [
            ['Opening Cash Balance',            fmt(f['wc_loan']*0.10), fmt(f['wc_loan']*0.15), fmt(f['wc_loan']*0.20)],
            ['Cash from Operations (PAT+Dep)',  fmt(f['pat'][1]+dep),   fmt(f['pat'][2]+dep),   fmt(f['pat'][3]+dep)],
            ['Less: Term Loan Repayment',       fmt(principal),         fmt(principal),         fmt(principal)],
            ['Less: Capex / Investments',       fmt(f['plant']*0.05),   fmt(f['plant']*0.03),   fmt(f['plant']*0.02)],
            ['Net Cash Flow',                   fmt(f['pat'][1]+dep-principal-f['plant']*0.05), fmt(f['pat'][2]+dep-principal-f['plant']*0.03), fmt(f['pat'][3]+dep-principal-f['plant']*0.02)],
        ]
    elif index == 10:
        dep = f['depreciation']
        fa1 = f['plant'] + f['building'] + f['misc_fixed'] - dep
        fa2 = fa1 - dep
        fa3 = fa2 - dep
        headers = ['Item', 'Year 1', 'Year 2', 'Year 3']
        rows = [
            ['Fixed Assets (Net Block)',  fmt(fa1),                                       fmt(fa2),                                       fmt(fa3)],
            ['Current Assets',           fmt(f['wc_margin']+f['wc_loan']),               fmt(f['wc_margin']+f['wc_loan']*0.9),           fmt(f['wc_margin']+f['wc_loan']*0.8)],
            ['Total Assets',             fmt(fa1+f['wc_margin']+f['wc_loan']),           fmt(fa2+f['wc_margin']+f['wc_loan']*0.9),      fmt(fa3+f['wc_margin']+f['wc_loan']*0.8)],
            ['Equity + Reserves',        fmt(f['equity']+f['pat'][1]),                   fmt(f['equity']+f['pat'][1]+f['pat'][2]),       fmt(f['equity']+sum(f['pat'][y] for y in [1,2,3]))],
            ['Term Loan (Outstanding)',  fmt(f['term_loan']*0.85),                       fmt(f['term_loan']*0.70),                       fmt(f['term_loan']*0.55)],
            ['Working Capital Loan',     fmt(f['wc_loan']),                              fmt(f['wc_loan']*0.9),                          fmt(f['wc_loan']*0.8)],
        ]
    elif index == 11:
        headers = ['Item', 'Year 1', 'Year 2', 'Year 3']
        dscr = {y: f['ebitda'][y] / (f['interest'][y] + principal) for y in [1, 2, 3]}
        rows = [
            ['Opening Loan Balance',    fmt(f['term_loan']),        fmt(f['term_loan']*0.85), fmt(f['term_loan']*0.70)],
            ['Interest @ 12% p.a.',     fmt(f['interest'][1]),      fmt(f['interest'][2]),    fmt(f['interest'][3])],
            ['Principal Repayment',     fmt(principal),             fmt(principal),           fmt(principal)],
            ['Total Debt Service',      fmt(f['interest'][1]+principal), fmt(f['interest'][2]+principal), fmt(f['interest'][3]+principal)],
            ['DSCR (min. 1.25x)',       f'{dscr[1]:.2f}x',         f'{dscr[2]:.2f}x',        f'{dscr[3]:.2f}x'],
            ['Closing Loan Balance',    fmt(f['term_loan']*0.85),   fmt(f['term_loan']*0.70), fmt(f['term_loan']*0.55)],
        ]
    elif index == 12:
        fc = f['depreciation'] + f['interest'][2] + f['labor'][2]*0.5 + f['admin'][2]
        vc_ratio = (f['rm'][2] + f['util'][2] + f['labor'][2]*0.5) / f['rev'][2]
        cm = 1 - vc_ratio
        bep = fc / cm if cm > 0 else 0
        headers = ['Parameter', 'Value', 'Notes', '']
        rows = [
            ['Fixed Costs (Annual)',      fmt(fc),                    'Depreciation + Interest + 50% Labour + Admin', ''],
            ['Variable Cost Ratio',       f'{vc_ratio*100:.1f}%',    'RM + Utilities + 50% Labour / Revenue', ''],
            ['Contribution Margin (%)',   f'{cm*100:.1f}%',          'Revenue − Variable Costs / Revenue', ''],
            ['Break-even Revenue',        fmt(bep),                   'Fixed Costs / Contribution Margin', ''],
            ['Break-even (% Capacity)',   f'{bep/f["rev"][3]*100:.1f}%', 'BEP Revenue / Full-capacity Revenue', ''],
        ]
    elif index == 13:
        d1 = f['rev'][1] / 365
        d2 = f['rev'][2] / 365
        headers = ['Item', 'Days', 'Year 1 Amount', 'Year 2 Amount']
        rows = [
            ['Raw Material Holding',  '30 days', fmt(d1*30*0.40), fmt(d2*30*0.40)],
            ['WIP',                   '7 days',  fmt(d1*7*0.55),  fmt(d2*7*0.55)],
            ['Finished Goods',        '15 days', fmt(d1*15*0.65), fmt(d2*15*0.65)],
            ['Debtors (Receivables)', '45 days', fmt(d1*45),      fmt(d2*45)],
            ['Less: Creditors',       '30 days', fmt(d1*30*0.40), fmt(d2*30*0.40)],
            ['Net Working Capital',   '67 days', fmt(f['wc_margin']), fmt(f['wc_margin']*1.1)],
        ]
    elif index == 14:
        base_dscr = f['ebitda'][2] / (f['interest'][2] + principal)
        headers = ['Scenario', 'Revenue', 'PAT', 'DSCR']
        rows = [
            ['Base Case (75% Util.)',       '100%',  fmt(f['pat'][2]),         f'{base_dscr:.2f}x'],
            ['Optimistic (+15% Revenue)',   '+15%',  fmt(f['pat'][2]*1.35),    f'{base_dscr*1.20:.2f}x'],
            ['Pessimistic (−15% Revenue)',  '−15%',  fmt(f['pat'][2]*0.55),    f'{base_dscr*0.75:.2f}x'],
            ['High Input Cost (+10%)',      'Base',  fmt(f['pat'][2]*0.80),    f'{base_dscr*0.90:.2f}x'],
            ['Low Capacity (50% Util.)',    '−33%',  fmt(f['pat'][2]*0.35),    f'{base_dscr*0.60:.2f}x'],
        ]
    else:  # index == 15 - Financial Ratios
        gm = {y: (f['rev'][y] - f['rm'][y] - f['util'][y]) / f['rev'][y] for y in [1, 2, 3]}
        nm = {y: f['pat'][y] / f['rev'][y] for y in [1, 2, 3]}
        roe = {y: f['pat'][y] / f['equity'] for y in [1, 2, 3]}
        roc = {y: f['ebit'][y] / f['total'] for y in [1, 2, 3]}
        dscr = {y: f['ebitda'][y] / (f['interest'][y] + principal) for y in [1, 2, 3]}
        headers = ['Ratio', 'Year 1', 'Year 2', 'Year 3']
        rows = [
            ['Gross Margin (%)',         f'{gm[1]*100:.1f}%',   f'{gm[2]*100:.1f}%',   f'{gm[3]*100:.1f}%'],
            ['Net Profit Margin (%)',    f'{nm[1]*100:.1f}%',   f'{nm[2]*100:.1f}%',   f'{nm[3]*100:.1f}%'],
            ['Return on Equity (%)',     f'{roe[1]*100:.1f}%',  f'{roe[2]*100:.1f}%',  f'{roe[3]*100:.1f}%'],
            ['Return on Capital (%)',    f'{roc[1]*100:.1f}%',  f'{roc[2]*100:.1f}%',  f'{roc[3]*100:.1f}%'],
            ['DSCR',                    f'{dscr[1]:.2f}x',     f'{dscr[2]:.2f}x',     f'{dscr[3]:.2f}x'],
            ['Current Ratio',           '1.50x',               '1.80x',               '2.10x'],
        ]
    return headers, rows


def add_financial_table_pack(doc: Document, submission: Dict[str, Any]) -> None:
    """
    Add a 15-table financial pack inside Chapter 6.
    """
    table_titles = [
        "Table 6.1 - Project Cost Breakup",
        "Table 6.2 - Means of Finance",
        "Table 6.3 - Installed Capacity and Utilization",
        "Table 6.4 - Revenue Projection",
        "Table 6.5 - Raw Material Cost Estimate",
        "Table 6.6 - Utility and Operating Cost Estimate",
        "Table 6.7 - Employee and Overhead Cost",
        "Table 6.8 - Profit and Loss Projection",
        "Table 6.9 - Cash Flow Projection",
        "Table 6.10 - Balance Sheet Projection",
        "Table 6.11 - Debt Service Schedule",
        "Table 6.12 - Break-even Analysis",
        "Table 6.13 - Working Capital Cycle",
        "Table 6.14 - Sensitivity Analysis",
        "Table 6.15 - Financial Ratios and KPIs",
    ]

    budget_raw = submission.get('budget', submission.get('total_investment', '5000000'))
    budget_inr = _parse_budget_inr(str(budget_raw))
    f = _compute_financials(budget_inr)

    doc.add_heading('6.1 Financial Tables (Target: 15 Pages)', level=2)
    doc.add_paragraph(
        "The following financial table pack presents a 3-year financial projection derived "
        f"from the total project cost of {_fmt(budget_inr)}. All figures use standard "
        "industry assumptions and should be validated against actual vendor quotes."
    )

    for index, title in enumerate(table_titles, start=1):
        heading_para = doc.add_paragraph(title)
        if heading_para.runs:
            heading_para.runs[0].bold = True

        headers, rows = _build_table_data(index, f)
        num_cols = len(headers)
        table = doc.add_table(rows=1 + len(rows), cols=num_cols)
        table.style = 'Table Grid'

        # Header row
        hdr_cells = table.rows[0].cells
        for col_index, header in enumerate(headers):
            hdr_cells[col_index].text = header
            for para in hdr_cells[col_index].paragraphs:
                for run in para.runs:
                    run.bold = True

        # Data rows
        for row_index, row_data in enumerate(rows, start=1):
            for col_index, cell_value in enumerate(row_data):
                table.rows[row_index].cells[col_index].text = str(cell_value)

        doc.add_paragraph()


def get_or_generate_section(
    submission_id: int, 
    section_name: str, 
    submission_data: Dict[str, Any],
    force: bool = False,
    generation_mode: str = "plain",
    max_tokens: int = 1600,
    extra_context: Dict[str, Any] = None,
) -> str:
    """
    Get cached section or generate new one using LLM.
    
    Args:
        submission_id: The submission ID
        section_name: Name of the section to generate
        submission_data: Dictionary containing submission data
        force: If True, regenerate even if cached
        
    Returns:
        Generated or cached section content
    """
    # Check cache unless force is True
    if not force:
        cached = get_cached_section(submission_id, section_name)
        if cached:
            return cached
    
    # Generate new content
    rendered_prompt = get_section_prompt(section_name, submission_data, extra_context)
    content = llm_client.generate(rendered_prompt, max_tokens=max_tokens, mode=generation_mode)
    
    # Cache the result
    save_section(submission_id, section_name, content)
    
    return content


def build_doc(submission: Dict[str, Any], submission_id: int, force: bool = False) -> bytes:
    """
    Build a Word document from submission data with AI-generated content.
    
    Args:
        submission: Dictionary containing submission data
        submission_id: The submission ID for caching
        force: If True, regenerate sections even if cached
        
    Returns:
        Bytes of the generated .docx file
    """
    # Identify missing inputs for assumptions
    missing_inputs = identify_missing_inputs(submission)
    
    # Add missing_inputs info to submission data for prompt rendering
    submission_with_context = {**submission}
    if missing_inputs:
        submission_with_context['missing_inputs'] = ', '.join(missing_inputs)
    else:
        submission_with_context['missing_inputs'] = 'None'
    
    section_names = [
        'introduction',
        'regulatory_framework',
        'market_assessment',
        'business_operating_model',
        'equipment_profiles',
        'financial_feasibility',
        'risk_assessment',
        'caveats',
        'appendices',
    ]

    # Split into two rounds to stay within Anthropic TPM limits.
    # Round 1: chapters 2-4 (introduction → market_assessment)
    # Round 2: chapters 5-8 + appendices (business_operating_model → appendices)
    # Executive Summary is generated LAST so it can reference all other sections.
    round_1 = section_names[:3]
    round_2 = section_names[3:]

    total_calls = len(section_names) + 1  # +1 for executive_summary call
    section_content: Dict[str, str] = {}

    def _generate_sections(batch, offset):
        for i, section_name in enumerate(batch):
            generation_mode = Config.resolve_section_mode(section_name)
            section_max_tokens = Config.WEB_SECTION_MAX_TOKENS if generation_mode == "web" else Config.PLAIN_SECTION_MAX_TOKENS
            upsert_report_status(
                submission_id, "generating",
                sections_done=offset + i + 1,
                sections_total=total_calls,
                current_section=f"{SECTION_LABELS.get(section_name, section_name)} ({generation_mode})",
            )
            section_content[section_name] = get_or_generate_section(
                submission_id,
                section_name,
                submission_with_context,
                force,
                generation_mode,
                section_max_tokens,
            )

    # Round 1
    _generate_sections(round_1, offset=0)

    # Round 2 (no cooldown — web search is disabled so TPM pressure is minimal)
    _generate_sections(round_2, offset=len(round_1))

    # Build financial highlights for executive summary context
    budget_raw = submission.get('budget', submission.get('total_investment', '5000000'))
    budget_inr = _parse_budget_inr(str(budget_raw))
    fin = _compute_financials(budget_inr)
    gm_pct = {y: (fin['rev'][y] - fin['rm'][y] - fin['util'][y]) / fin['rev'][y] * 100 for y in [1, 2, 3]}
    nm_pct = {y: fin['pat'][y] / fin['rev'][y] * 100 for y in [1, 2, 3]}
    principal = fin['term_loan'] * 0.15
    dscr = {y: fin['ebitda'][y] / (fin['interest'][y] + principal) for y in [1, 2, 3]}
    payback_yr = next((y for y in [1, 2, 3] if fin['pat'][y] > 0), 3)
    financial_highlights = (
        f"- Total Project Cost: {_fmt(budget_inr)}\n"
        f"- Gross Margin: Y1 {gm_pct[1]:.1f}% / Y2 {gm_pct[2]:.1f}% / Y3 {gm_pct[3]:.1f}%\n"
        f"- PAT Margin: Y1 {nm_pct[1]:.1f}% / Y2 {nm_pct[2]:.1f}% / Y3 {nm_pct[3]:.1f}%\n"
        f"- EBITDA: Y1 {_fmt(fin['ebitda'][1])} / Y2 {_fmt(fin['ebitda'][2])} / Y3 {_fmt(fin['ebitda'][3])}\n"
        f"- DSCR: Y1 {dscr[1]:.2f}x / Y2 {dscr[2]:.2f}x / Y3 {dscr[3]:.2f}x (min DSCR: {min(dscr.values()):.2f}x)\n"
        f"- Indicative Payback: Year {payback_yr}\n"
        f"- Debt: {_fmt(fin['term_loan'])} ({submission.get('debt_percentage', 50)}% of project cost)\n"
        f"- Equity: {_fmt(fin['equity'])} ({submission.get('equity_percentage', 30)}% of project cost)"
    )

    # Extract brief context snippets from completed sections
    risk_context = (section_content.get('risk_assessment', '') or '')[:600].strip()
    market_context = (section_content.get('market_assessment', '') or '')[:400].strip()

    # Generate Executive Summary LAST with full context
    upsert_report_status(
        submission_id, "generating",
        sections_done=total_calls,
        sections_total=total_calls,
        current_section="Executive Summary (final step)",
    )
    section_content['executive_summary'] = get_or_generate_section(
        submission_id,
        'executive_summary',
        submission_with_context,
        force,
        Config.resolve_section_mode('executive_summary'),
        Config.PLAIN_SECTION_MAX_TOKENS,
        extra_context={
            'financial_highlights': financial_highlights,
            'risk_context': risk_context or 'Risk assessment not yet available.',
            'market_context': market_context or 'Market assessment not yet available.',
        },
    )

    # Validate and sanitize links in all generated sections before rendering output.
    upsert_report_status(
        submission_id, "generating",
        sections_done=total_calls,
        sections_total=total_calls,
        current_section="Validating source links",
    )
    link_validation_cache: Dict[str, bool] = {}
    for key in list(section_content.keys()):
        section_content[key] = _sanitize_invalid_links(section_content[key], link_validation_cache)

    # Mark financial tables as the final step
    upsert_report_status(
        submission_id, "generating",
        sections_done=total_calls,
        sections_total=total_calls,
        current_section="Finalizing financial tables",
    )

    doc = Document()

    apply_report_formatting(doc)
    
    # Title
    title = doc.add_heading('Project Feasibility Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle with project name
    subtitle = doc.add_paragraph(f"Project: {submission.get('business_idea', 'N/A')}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.size = Pt(12)
    subtitle_format.font.italic = True
    
    doc.add_paragraph()  # Spacing

    # Table of Contents (Word will populate on open)
    _add_toc(doc)

    # Chapters 1-8 (target 90 pages in aggregate per output specification)
    doc.add_page_break()
    doc.add_heading('Chapter 1: Executive Summary (Target: 2 Pages)', level=1)
    render_markdown_to_doc(doc, section_content['executive_summary'])

    doc.add_page_break()
    doc.add_heading('Chapter 2: Introduction (Target: 6 Pages)', level=1)
    render_markdown_to_doc(doc, section_content['introduction'])

    doc.add_page_break()
    doc.add_heading('Chapter 3: Regulatory Framework (Target: 10 Pages)', level=1)
    render_markdown_to_doc(doc, section_content['regulatory_framework'])

    doc.add_page_break()
    doc.add_heading('Chapter 4: Market Assessment (Target: 16 Pages)', level=1)
    render_markdown_to_doc(doc, section_content['market_assessment'])

    doc.add_page_break()
    doc.add_heading('Chapter 5: Business and Operating Model (Target: 23 Pages)', level=1)
    render_markdown_to_doc(doc, section_content['business_operating_model'])
    doc.add_heading('5.1 Illustrated Key Equipment Profiles (Target: 5-6 Pages)', level=2)
    render_markdown_to_doc(doc, section_content['equipment_profiles'])

    doc.add_page_break()
    doc.add_heading('Chapter 6: Financial Feasibility (Target: 24 Pages)', level=1)
    render_markdown_to_doc(doc, section_content['financial_feasibility'])
    add_financial_table_pack(doc, submission)

    doc.add_page_break()
    doc.add_heading('Chapter 7: Risk Assessment & Mitigation (Target: 6 Pages)', level=1)
    render_markdown_to_doc(doc, section_content['risk_assessment'])

    doc.add_page_break()
    doc.add_heading('Chapter 8: Caveats (Target: 3 Pages)', level=1)
    render_markdown_to_doc(doc, section_content['caveats'])

    # Additional chapterized context sections
    doc.add_heading('Project Timeline', level=1)
    timeline_para = doc.add_paragraph()
    timeline_para.add_run(f"Start Date: ").bold = True
    timeline_para.add_run(f"{submission.get('start_date', 'N/A')}\n")
    timeline_para.add_run(f"Target Launch Date: ").bold = True
    timeline_para.add_run(f"{submission.get('target_launch_date', 'N/A')}")

    doc.add_heading('Budget Overview', level=1)
    budget_para = doc.add_paragraph()
    budget_para.add_run(f"Total Project Budget: ").bold = True
    budget_para.add_run(f"{submission.get('budget', 'N/A')} currency units")

    # Appendices are generated but treated outside the 90-page chapter count.
    doc.add_page_break()
    doc.add_heading('Appendices', level=1)
    render_markdown_to_doc(doc, section_content['appendices'])
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph("---")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_text = doc.add_paragraph("This report was generated automatically by the Project Report Automation system.")
    footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_text.runs[0].font.size = Pt(9)
    footer_text.runs[0].font.italic = True
    
    # Save to bytes
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()
